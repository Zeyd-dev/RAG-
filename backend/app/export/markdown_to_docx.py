"""
Turns a chat answer (Markdown text with [S1]/[S2]-style citation tokens,
as produced by the Groq provider) into a downloadable .docx file --
this is what powers the "Export as Word" button so a report-style answer
can leave the app as an actual file, not just chat text.

Deliberately a small hand-rolled Markdown reader rather than a new
dependency: the model only ever produces a constrained subset of Markdown
(headings, bold, bullet/numbered lists, GFM tables, plain paragraphs) per
the system prompt in llm/groq_provider.py, so a full CommonMark parser
would be more machinery than the job needs. If richer Markdown shows up
later, extend the block/inline handling here rather than reaching for a
general parser -- keeps this module dependency-free (python-docx is
already a requirement for DOCX ingestion).
"""
import io
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

CITATION_TOKEN = re.compile(r"\[S(\d+)\]")
BOLD_OR_CITATION = re.compile(r"(\*\*.+?\*\*)|(\[S\d+\])")
TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
TABLE_SEPARATOR = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")
ORDERED_ITEM = re.compile(r"^\s*\d+[.)]\s+(.*)$")
UNORDERED_ITEM = re.compile(r"^\s*[-*]\s+(.*)$")
HEADING = re.compile(r"^(#{1,3})\s+(.*)$")


def _add_inline_runs(paragraph, text: str) -> None:
    """Write `text` into `paragraph`, honoring **bold** spans and turning
    [S<n>] citation tokens into small superscript numbers -- mirrors what
    the frontend's CitationChip does, just as plain-text superscript since
    a Word doc has no click handler to hang a tooltip off of."""
    pos = 0
    for match in BOLD_OR_CITATION.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        bold_span, citation_span = match.groups()
        if bold_span:
            run = paragraph.add_run(bold_span[2:-2])
            run.bold = True
        elif citation_span:
            num = CITATION_TOKEN.match(citation_span).group(1)
            run = paragraph.add_run(num)
            run.font.superscript = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Reads a GFM table starting at lines[start] (header row), lines[start+1]
    is the --- separator row, and subsequent lines are data rows. Returns
    the parsed rows (including header) and the index just past the table."""
    rows = []
    i = start
    while i < len(lines) and TABLE_ROW.match(lines[i]):
        if i == start + 1 and TABLE_SEPARATOR.match(lines[i]) and not lines[i].strip().strip("|-: "):
            i += 1
            continue
        if i == start + 1 and re.fullmatch(r"[\s|:-]+", lines[i].strip()):
            i += 1
            continue
        cells = [c.strip() for c in TABLE_ROW.match(lines[i]).group(1).split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def render_markdown_to_docx(
    content: str,
    citations: Optional[list[dict]] = None,
    title: Optional[str] = None,
) -> io.BytesIO:
    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    lines = content.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        heading_match = HEADING.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2).strip(), level=level)
            i += 1
            continue

        if TABLE_ROW.match(line) and i + 1 < len(lines) and re.fullmatch(
            r"[\s|:-]+", lines[i + 1].strip()
        ):
            rows, i = _parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row_cells in enumerate(rows):
                    for c, cell_text in enumerate(row_cells):
                        if c >= len(table.columns):
                            continue
                        cell_paragraph = table.cell(r, c).paragraphs[0]
                        _add_inline_runs(cell_paragraph, cell_text)
                        if r == 0:
                            for run in cell_paragraph.runs:
                                run.bold = True
                doc.add_paragraph()
            continue

        ordered_match = ORDERED_ITEM.match(line)
        if ordered_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, ordered_match.group(1))
            i += 1
            continue

        unordered_match = UNORDERED_ITEM.match(line)
        if unordered_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, unordered_match.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, line.strip())
        i += 1

    if citations:
        doc.add_heading("Sources", level=2)
        for idx, c in enumerate(citations, start=1):
            p = doc.add_paragraph(style="List Number")
            page = c.get("page")
            label = f"{c.get('filename', 'unknown')}" + (f", page {page}" if page is not None else "")
            run = p.add_run(label)
            run.italic = True

    footer_note = doc.sections[0].footer.paragraphs[0]
    footer_note.text = "Generated by RAG NoteBook -- verify against the original source documents."
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_note.runs:
        run.font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
